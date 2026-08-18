#!/usr/bin/env python3
"""Generate Word (.docx) requirement templates from static Markdown sources."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError as exc:
    raise SystemExit(
        "python-docx is required. Run: python3 -m venv .venv-docx && "
        ".venv-docx/bin/pip install python-docx"
    ) from exc

ROOT = Path(__file__).resolve().parents[1]
STATIC = ROOT / "static" / "requirement-templates"


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    table_rows: list[str] = []
    i = 0

    def flush_table() -> None:
        nonlocal table_rows
        if not table_rows:
            return
        cols = max(len(r.split("|")) - 2 for r in table_rows if "|" in r)
        cols = max(cols, 1)
        table = doc.add_table(rows=0, cols=cols)
        table.style = "Table Grid"
        for row_idx, row_line in enumerate(table_rows):
            cells = [c.strip() for c in row_line.strip("|").split("|")]
            while len(cells) < cols:
                cells.append("")
            if row_idx == 1 and all(re.match(r"^[-:\s]+$", c) for c in cells):
                continue
            row_cells = table.add_row().cells
            for ci, val in enumerate(cells):
                row_cells[ci].text = val
        doc.add_paragraph()
        table_rows = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("|") and "|" in line[1:]:
            table_rows.append(line)
            i += 1
            continue
        if table_rows:
            flush_table()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            doc.add_paragraph()
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith(":::") or stripped.startswith("```"):
            i += 1
            continue
        elif stripped.startswith("- [ ]"):
            doc.add_paragraph("☐ " + stripped[5:].strip(), style="List Bullet")
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        else:
            paragraph = doc.add_paragraph()
            for part in re.split(r"(\*\*.*?\*\*)", stripped):
                if part.startswith("**") and part.endswith("**"):
                    run = paragraph.add_run(part[2:-2])
                    run.bold = True
                else:
                    paragraph.add_run(part)
        i += 1

    if table_rows:
        flush_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def main() -> int:
    md_files = sorted(STATIC.glob("*-template.md"))
    if not md_files:
        print(f"No templates found in {STATIC}", file=sys.stderr)
        return 1
    for md_path in md_files:
        docx_path = md_path.with_suffix(".docx")
        md_to_docx(md_path, docx_path)
        print(f"Created {docx_path.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
